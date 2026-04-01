"""Word document generation service."""

from __future__ import annotations

import io
import logging
from typing import List, Optional, Tuple

from lxml import etree

from app.services.omml import OMML_NS, latex_to_omml

logger = logging.getLogger(__name__)

# Content block types
TEXT_BLOCK = "text"
FORMULA_BLOCK = "formula"


def _add_equation_paragraph(doc, latex: str) -> bool:
    """Insert an editable OMML equation as a new paragraph in *doc*.

    The equation is embedded as ``<m:oMathPara><m:oMath>…</m:oMath></m:oMathPara>``
    inside a Word paragraph, making it fully editable in Word's equation
    editor and in MathType.

    Parameters
    ----------
    doc:
        A ``python-docx`` ``Document`` instance.
    latex:
        LaTeX math string (without surrounding ``$`` signs).

    Returns
    -------
    ``True`` if the equation was inserted successfully, ``False`` otherwise.
    """
    oMath = latex_to_omml(latex)
    if oMath is None:
        logger.warning("Could not convert LaTeX to OMML: %s", latex)
        return False

    para = doc.add_paragraph()
    p = para._p

    # Remove any empty <w:r> run that python-docx adds automatically
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for run in p.findall(f"{{{W_NS}}}r"):
        p.remove(run)

    # Wrap the <m:oMath> element in <m:oMathPara>
    oMathPara = etree.SubElement(p, f"{{{OMML_NS}}}oMathPara")
    oMath_dest = etree.SubElement(oMathPara, f"{{{OMML_NS}}}oMath")
    for child in list(oMath):
        oMath_dest.append(child)

    return True


class WordGeneratorService:
    """Assembles text and formula blocks into a .docx document."""

    def generate(
        self,
        blocks: List[Tuple[str, str]],
        title: Optional[str] = None,
    ) -> bytes:
        """Build a Word document from *blocks* and return the raw bytes.

        Parameters
        ----------
        blocks:
            A list of ``(block_type, content)`` tuples where *block_type* is
            either ``"text"`` or ``"formula"``, and *content* is the
            recognised text or LaTeX string.
        title:
            Optional document title inserted as a heading.

        Returns
        -------
        Raw ``.docx`` bytes.
        """
        from docx import Document  # noqa: PLC0415
        from docx.shared import Pt  # noqa: PLC0415

        doc = Document()

        if title:
            heading = doc.add_heading(title, level=1)
            heading.runs[0].font.size = Pt(16)

        for block_type, content in blocks:
            if block_type == TEXT_BLOCK:
                doc.add_paragraph(content)
            elif block_type == FORMULA_BLOCK:
                inserted = _add_equation_paragraph(doc, content)
                if not inserted:
                    # Fallback: insert the raw LaTeX as plain text
                    doc.add_paragraph(f"[Formula: {content}]")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
